"""
报告生成服务
支持：
- generate_weekly_report(project_id, db) -> Word (.docx)
- generate_monthly_report(project_id, year, month, db) -> Word (.docx)
- generate_issue_risk_excel(project_id, db) -> Excel (.xlsx)
- generate_manday_excel(project_id, db) -> Excel (.xlsx)
- generate_status_excel(db) -> Excel (.xlsx) 多项目一览
"""

import io
from datetime import date, datetime
from typing import Optional
from sqlalchemy.orm import Session
from sqlalchemy import func

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from app.models.project import Project
from app.models.milestone import Milestone, Task
from app.models.issue import Issue
from app.models.risk import Risk
from app.models.manday import ManDay

# ──────────────────────────────────────────────────────────
# 颜色常量
# ──────────────────────────────────────────────────────────
COLOR_HEADER = "2E4057"   # 深蓝-灰
COLOR_SUBHEADER = "4472C4"
COLOR_GREEN = "70AD47"
COLOR_RED = "FF0000"
COLOR_ORANGE = "FFC000"
COLOR_YELLOW = "FFFF00"

STATUS_COLOR = {
    "正常": COLOR_GREEN,
    "预警": COLOR_ORANGE,
    "延期": COLOR_RED,
    "暂停": "808080",
    "已完成": COLOR_GREEN,
}

RISK_COLOR = {
    "极高": COLOR_RED,
    "高": "FF4444",
    "中": COLOR_ORANGE,
    "低": COLOR_GREEN,
    "极低": "A9D18E",
}

SEVERITY_COLOR = {"高": COLOR_RED, "中": COLOR_ORANGE, "低": COLOR_GREEN}


# ──────────────────────────────────────────────────────────
# Word 辅助函数
# ──────────────────────────────────────────────────────────
def _fmt_date(d) -> str:
    if d is None:
        return "-"
    if isinstance(d, (date, datetime)):
        return d.strftime("%Y-%m-%d")
    return str(d)


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0x57)


def _add_table(doc: Document, headers: list, rows: list):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # 表头行
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(
            _make_cell_color(COLOR_HEADER)
        )
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val) if val is not None else "-"
    doc.add_paragraph()
    return table


def _make_cell_color(hex_color: str):
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    return shd


# ──────────────────────────────────────────────────────────
# Excel 辅助函数
# ──────────────────────────────────────────────────────────
def _excel_header_style(cell, bg_hex: str = COLOR_HEADER):
    cell.font = Font(bold=True, color="FFFFFF", size=11)
    cell.fill = PatternFill("solid", fgColor=bg_hex)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    _apply_border(cell)


def _apply_border(cell):
    thin = Side(style="thin")
    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)


def _data_style(cell, align="left"):
    cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
    _apply_border(cell)


def _color_cell(cell, hex_color: str):
    cell.fill = PatternFill("solid", fgColor=hex_color)


# ──────────────────────────────────────────────────────────
# 1. Word 周报
# ──────────────────────────────────────────────────────────
def generate_weekly_report(project_id: int, db: Session, report_date: Optional[date] = None) -> bytes:
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise ValueError("项目不存在")

    today = report_date or date.today()
    doc = Document()

    # 标题
    title = doc.add_heading("", 0)
    run = title.add_run(f"{project.name} — 项目周报")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表
    doc.add_paragraph()
    _add_heading(doc, "一、项目基本信息", level=2)
    info_table = doc.add_table(rows=5, cols=4)
    info_table.style = "Table Grid"
    info = [
        ("项目名称", project.name, "项目编号", project.code),
        ("客户/甲方", project.client or "-", "项目经理", project.manager or "-"),
        ("项目阶段", project.phase, "项目状态", project.status),
        ("计划开始", _fmt_date(project.plan_start), "计划结束", _fmt_date(project.plan_end)),
        ("实际开始", _fmt_date(project.actual_start), "报告日期", _fmt_date(today)),
    ]
    for r_idx, (l1, v1, l2, v2) in enumerate(info):
        row = info_table.rows[r_idx]
        for c_idx, (txt, bold) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
            cell = row.cells[c_idx]
            cell.text = txt
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(txt)
            run.bold = bold
            if bold:
                cell._tc.get_or_add_tcPr().append(_make_cell_color("E8EDF3"))

    # 里程碑进度
    doc.add_paragraph()
    _add_heading(doc, "二、里程碑进度", level=2)
    milestones = db.query(Milestone).filter(Milestone.project_id == project_id).order_by(Milestone.order_index).all()
    if milestones:
        ms_rows = [
            (ms.name, ms.status, _fmt_date(ms.plan_date), _fmt_date(ms.actual_date), ms.description or "-")
            for ms in milestones
        ]
        _add_table(doc, ["里程碑", "状态", "计划日期", "实际日期", "备注"], ms_rows)
    else:
        doc.add_paragraph("暂无里程碑记录。")

    # 本周任务
    doc.add_paragraph()
    _add_heading(doc, "三、本周工作情况", level=2)
    tasks = db.query(Task).filter(Task.project_id == project_id, Task.status.in_(["进行中", "已完成"])).all()
    if tasks:
        task_rows = [
            (t.name, t.assignee or "-", t.status, f"{t.progress}%", _fmt_date(t.plan_end), t.notes or "-")
            for t in tasks
        ]
        _add_table(doc, ["工作项", "负责人", "状态", "进度", "计划完成", "备注"], task_rows)
    else:
        doc.add_paragraph("暂无进行中工作项。")

    # 未开始任务 = 下周计划
    doc.add_paragraph()
    _add_heading(doc, "四、下周工作计划", level=2)
    next_tasks = db.query(Task).filter(Task.project_id == project_id, Task.status == "未开始").all()
    if next_tasks:
        nt_rows = [
            (t.name, t.assignee or "-", _fmt_date(t.plan_start), _fmt_date(t.plan_end))
            for t in next_tasks
        ]
        _add_table(doc, ["工作项", "负责人", "计划开始", "计划结束"], nt_rows)
    else:
        doc.add_paragraph("暂无待开始工作项。")

    # 问题
    doc.add_paragraph()
    _add_heading(doc, "五、问题跟踪", level=2)
    issues = db.query(Issue).filter(Issue.project_id == project_id, Issue.status != "已关闭").all()
    if issues:
        i_rows = [(i.title, i.severity, i.assignee or "-", _fmt_date(i.due_date), i.status) for i in issues]
        _add_table(doc, ["问题", "等级", "负责人", "期望解决日期", "状态"], i_rows)
    else:
        doc.add_paragraph("暂无未关闭问题。")

    # 风险
    doc.add_paragraph()
    _add_heading(doc, "六、风险跟踪", level=2)
    risks = db.query(Risk).filter(Risk.project_id == project_id, Risk.status == "开放").all()
    if risks:
        r_rows = [(r.title, r.probability, r.impact, r.level or "-", r.assignee or "-", r.mitigation or "-") for r in risks]
        _add_table(doc, ["风险", "概率", "影响", "等级", "负责人", "应对措施"], r_rows)
    else:
        doc.add_paragraph("暂无开放风险。")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────
# 2. Word 月报
# ──────────────────────────────────────────────────────────
def generate_monthly_report(project_id: int, year: int, month: int, db: Session) -> bytes:
    from calendar import month_name
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise ValueError("项目不存在")

    doc = Document()
    title = doc.add_heading("", 0)
    run = title.add_run(f"{project.name} — {year}年{month:02d}月项目月报")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 项目总览
    doc.add_paragraph()
    _add_heading(doc, "一、项目总览", level=2)
    used_mandays = db.query(func.sum(ManDay.days)).filter(ManDay.project_id == project_id).scalar() or 0
    overview = [
        ("项目名称", project.name),
        ("项目状态", project.status),
        ("项目阶段", project.phase),
        ("计划结束", _fmt_date(project.plan_end)),
        ("预算人天", f"{project.budget_mandays:.1f}"),
        ("已用人天", f"{used_mandays:.1f}"),
        ("剩余人天", f"{max(project.budget_mandays - used_mandays, 0):.1f}"),
    ]
    ov_table = doc.add_table(rows=len(overview), cols=2)
    ov_table.style = "Table Grid"
    for i, (k, v) in enumerate(overview):
        row = ov_table.rows[i]
        row.cells[0].text = k
        row.cells[0]._tc.get_or_add_tcPr().append(_make_cell_color("E8EDF3"))
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = v

    # 里程碑总结
    doc.add_paragraph()
    _add_heading(doc, "二、里程碑完成情况", level=2)
    milestones = db.query(Milestone).filter(Milestone.project_id == project_id).order_by(Milestone.order_index).all()
    if milestones:
        ms_rows = [
            (ms.name, ms.status, _fmt_date(ms.plan_date), _fmt_date(ms.actual_date))
            for ms in milestones
        ]
        _add_table(doc, ["里程碑", "状态", "计划日期", "实际日期"], ms_rows)

    # 本月人天统计
    doc.add_paragraph()
    _add_heading(doc, "三、本月人天投入", level=2)
    import calendar
    _, last_day = calendar.monthrange(year, month)
    month_start = date(year, month, 1)
    month_end = date(year, month, last_day)
    month_mds = db.query(ManDay).filter(
        ManDay.project_id == project_id,
        ManDay.work_date >= month_start,
        ManDay.work_date <= month_end,
    ).all()
    if month_mds:
        md_rows = [(m.staff_name, m.role or "-", _fmt_date(m.work_date), m.days, "是" if m.is_billable else "否", m.work_content or "-") for m in month_mds]
        _add_table(doc, ["人员", "角色", "日期", "人天", "计费", "工作内容"], md_rows)
        month_total = sum(m.days for m in month_mds)
        doc.add_paragraph(f"本月合计：{month_total:.1f} 人天")
    else:
        doc.add_paragraph("本月暂无人天记录。")

    # 未关闭问题与风险
    doc.add_paragraph()
    _add_heading(doc, "四、未关闭问题清单", level=2)
    issues = db.query(Issue).filter(Issue.project_id == project_id, Issue.status != "已关闭").all()
    if issues:
        i_rows = [(i.title, i.severity, i.assignee or "-", _fmt_date(i.due_date), i.status) for i in issues]
        _add_table(doc, ["问题", "等级", "负责人", "期望解决", "状态"], i_rows)
    else:
        doc.add_paragraph("暂无未关闭问题。")

    doc.add_paragraph()
    _add_heading(doc, "五、开放风险清单", level=2)
    risks = db.query(Risk).filter(Risk.project_id == project_id, Risk.status == "开放").all()
    if risks:
        r_rows = [(r.title, r.level or "-", r.assignee or "-", r.mitigation or "-") for r in risks]
        _add_table(doc, ["风险", "等级", "负责人", "应对措施"], r_rows)
    else:
        doc.add_paragraph("暂无开放风险。")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────
# 3. Excel 问题与风险台账
# ──────────────────────────────────────────────────────────
def generate_issue_risk_excel(project_id: int, db: Session) -> bytes:
    project = db.query(Project).filter(Project.id == project_id).first()
    wb = openpyxl.Workbook()

    # Sheet 1: Issues
    ws_i = wb.active
    ws_i.title = "问题台账"
    issue_headers = ["#", "问题标题", "描述", "严重等级", "来源", "负责人", "提出日期", "期望解决", "实际解决", "状态", "解决措施"]
    for c, h in enumerate(issue_headers, 1):
        cell = ws_i.cell(row=1, column=c, value=h)
        _excel_header_style(cell)

    issues = db.query(Issue).filter(Issue.project_id == project_id).order_by(Issue.raised_date.desc()).all()
    for r, issue in enumerate(issues, 2):
        row_data = [
            r - 1, issue.title, issue.description, issue.severity, issue.source,
            issue.assignee, _fmt_date(issue.raised_date), _fmt_date(issue.due_date),
            _fmt_date(issue.resolved_date), issue.status, issue.resolution,
        ]
        for c, val in enumerate(row_data, 1):
            cell = ws_i.cell(row=r, column=c, value=val or "")
            _data_style(cell)
            if c == 4 and issue.severity in SEVERITY_COLOR:  # 严重等级着色
                _color_cell(cell, SEVERITY_COLOR[issue.severity])

    _auto_col_width(ws_i)

    # Sheet 2: Risks
    ws_r = wb.create_sheet("风险台账")
    risk_headers = ["#", "风险标题", "描述", "概率", "影响", "风险等级", "负责人", "应对措施", "状态"]
    for c, h in enumerate(risk_headers, 1):
        cell = ws_r.cell(row=1, column=c, value=h)
        _excel_header_style(cell)

    risks = db.query(Risk).filter(Risk.project_id == project_id).all()
    for r, risk in enumerate(risks, 2):
        row_data = [
            r - 1, risk.title, risk.description, risk.probability, risk.impact,
            risk.level, risk.assignee, risk.mitigation, risk.status,
        ]
        for c, val in enumerate(row_data, 1):
            cell = ws_r.cell(row=r, column=c, value=val or "")
            _data_style(cell)
            if c == 6 and risk.level in RISK_COLOR:
                _color_cell(cell, RISK_COLOR[risk.level])

    _auto_col_width(ws_r)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────
# 4. Excel 人天统计报表
# ──────────────────────────────────────────────────────────
def generate_manday_excel(project_id: int, db: Session) -> bytes:
    project = db.query(Project).filter(Project.id == project_id).first()
    wb = openpyxl.Workbook()

    ws = wb.active
    ws.title = "人天明细"
    headers = ["#", "人员姓名", "角色", "工作日期", "投入人天", "是否计费", "工作内容"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        _excel_header_style(cell)

    mandays = db.query(ManDay).filter(ManDay.project_id == project_id).order_by(ManDay.work_date).all()
    for r, md in enumerate(mandays, 2):
        row_data = [r - 1, md.staff_name, md.role, _fmt_date(md.work_date), md.days, "是" if md.is_billable else "否", md.work_content]
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val or "")
            _data_style(cell)

    # 汇总 sheet
    ws2 = wb.create_sheet("人天汇总")
    ws2.cell(row=1, column=1, value=f"项目：{project.name if project else ''}")
    ws2.cell(row=2, column=1, value=f"预算人天：{project.budget_mandays if project else 0}")

    breakdown = (
        db.query(ManDay.staff_name, ManDay.role, func.sum(ManDay.days).label("total"),
                 func.sum(ManDay.days).filter(ManDay.is_billable == True).label("billable"))
        .filter(ManDay.project_id == project_id)
        .group_by(ManDay.staff_name, ManDay.role)
        .all()
    )
    sum_headers = ["人员姓名", "角色", "合计人天", "计费人天", "非计费人天"]
    for c, h in enumerate(sum_headers, 1):
        cell = ws2.cell(row=4, column=c, value=h)
        _excel_header_style(cell)
    for r, b in enumerate(breakdown, 5):
        total = b.total or 0
        billable = b.billable or 0
        row = [b.staff_name, b.role, round(total, 1), round(billable, 1), round(total - billable, 1)]
        for c, val in enumerate(row, 1):
            ws2.cell(row=r, column=c, value=val or "")

    _auto_col_width(ws)
    _auto_col_width(ws2)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────
# 5. Excel 多项目状态一览
# ──────────────────────────────────────────────────────────
def generate_status_excel(db: Session) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "项目状态一览"
    headers = ["项目编号", "项目名称", "客户", "项目经理", "阶段", "状态", "计划结束", "里程碑数", "未关闭问题", "开放风险", "已用人天", "预算人天"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        _excel_header_style(cell)

    projects = db.query(Project).all()
    for r, p in enumerate(projects, 2):
        used = db.query(func.sum(ManDay.days)).filter(ManDay.project_id == p.id).scalar() or 0
        open_issues = db.query(func.count(Issue.id)).filter(Issue.project_id == p.id, Issue.status != "已关闭").scalar() or 0
        open_risks = db.query(func.count(Risk.id)).filter(Risk.project_id == p.id, Risk.status == "开放").scalar() or 0
        ms_count = db.query(func.count(Milestone.id)).filter(Milestone.project_id == p.id).scalar() or 0

        row_data = [p.code, p.name, p.client, p.manager, p.phase, p.status,
                    _fmt_date(p.plan_end), ms_count, open_issues, open_risks,
                    round(used, 1), p.budget_mandays]
        for c, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=c, value=val or "")
            _data_style(cell, align="center")
            if c == 6 and p.status in STATUS_COLOR:  # 状态着色
                _color_cell(cell, STATUS_COLOR[p.status])

    _auto_col_width(ws)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────
# 工具：自动列宽
# ──────────────────────────────────────────────────────────
def _auto_col_width(ws):
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                val_len = len(str(cell.value)) if cell.value else 0
                # 中文字符按2个宽度计算
                cn_len = sum(2 if ord(c) > 127 else 1 for c in str(cell.value or ""))
                max_len = max(max_len, cn_len)
            except:
                pass
        ws.column_dimensions[col_letter].width = min(max_len + 4, 50)


# ──────────────────────────────────────────────────────────
# 6. Word 项目组合整体报告（Portfolio Report）
# ──────────────────────────────────────────────────────────
def generate_portfolio_report_word(db: Session, report_date: Optional[date] = None) -> bytes:
    """
    生成覆盖所有项目的组合整体报告（Word格式）
    包含5个部分：
      1. 执行摘要
      2. 重点关注项目（预警/延期）
      3. 全局重大风险与问题（高级别）
      4. 近期里程碑跟踪
      5. 项目状态一览表
    """
    today = report_date or date.today()
    from datetime import timedelta

    doc = Document()

    # ---- 封面标题 ----
    title = doc.add_heading("", 0)
    run = title.add_run("PMO 项目组合综合报告")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"报告日期：{_fmt_date(today)}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ─────────────────────────────────────────────
    # 第一部分：执行摘要
    # ─────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "一、执行摘要", level=2)

    all_projects = db.query(Project).all()
    total = len(all_projects)
    status_count = {}
    for p in all_projects:
        status_count[p.status] = status_count.get(p.status, 0) + 1

    in_progress = sum(status_count.get(s, 0) for s in ["正常", "预警", "延期"])
    completed = status_count.get("已完成", 0)
    paused = status_count.get("暂停", 0)
    attention = status_count.get("延期", 0) + status_count.get("预警", 0)

    total_budget = sum(p.budget_mandays for p in all_projects)
    total_used = db.query(func.sum(ManDay.days)).scalar() or 0
    total_issues = db.query(func.count(Issue.id)).filter(Issue.status != "已关闭").scalar() or 0
    total_risks = db.query(func.count(Risk.id)).filter(Risk.status == "开放").scalar() or 0

    summary_rows = [
        ("项目总数", str(total), "进行中项目", str(in_progress)),
        ("已完成项目", str(completed), "暂停项目", str(paused)),
        ("需重点关注（预警+延期）", str(attention), "未关闭高危问题", str(
            db.query(func.count(Issue.id)).filter(Issue.status != "已关闭", Issue.severity == "高").scalar() or 0
        )),
        ("全局预算人天", f"{total_budget:.1f}", "已消耗人天（合计）", f"{total_used:.1f}"),
    ]

    summary_table = doc.add_table(rows=len(summary_rows), cols=4)
    summary_table.style = "Table Grid"
    for r_idx, (l1, v1, l2, v2) in enumerate(summary_rows):
        row = summary_table.rows[r_idx]
        for c_idx, (txt, bold) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
            cell = row.cells[c_idx]
            cell.text = txt
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = bold
            else:
                run_obj = cell.paragraphs[0].add_run(txt)
                run_obj.bold = bold
            if bold:
                cell._tc.get_or_add_tcPr().append(_make_cell_color("D9E1F2"))

    doc.add_paragraph()

    # ─────────────────────────────────────────────
    # 第二部分：重点关注项目（预警 / 延期）
    # ─────────────────────────────────────────────
    _add_heading(doc, "二、重点关注项目（预警 / 延期）", level=2)
    attention_projects = [p for p in all_projects if p.status in ["预警", "延期"]]
    if attention_projects:
        for p in attention_projects:
            # 项目名称段
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(f"▶ [{p.status}] {p.name}（{p.code}）")
            proj_run.bold = True
            proj_run.font.color.rgb = RGBColor(0xFF, 0x40, 0x00) if p.status == "延期" else RGBColor(0xFF, 0xC0, 0x00)

            # 项目基本信息
            detail_rows = [
                (p.code, p.name, p.manager or "-", p.phase, p.status,
                 _fmt_date(p.plan_end), _fmt_date(p.actual_end) if p.actual_end else "进行中")
            ]
            # 该项目高级问题
            proj_issues = db.query(Issue).filter(
                Issue.project_id == p.id,
                Issue.severity == "高",
                Issue.status != "已关闭"
            ).all()
            if proj_issues:
                issue_rows = [(i.title, i.severity, i.assignee or "-", _fmt_date(i.due_date), i.status) for i in proj_issues]
                doc.add_paragraph("  高危未关闭问题：").runs[0].italic = True
                _add_table(doc, ["问题标题", "严重等级", "负责人", "期望解决", "状态"], issue_rows)
            # 高风险
            proj_risks = db.query(Risk).filter(
                Risk.project_id == p.id,
                Risk.impact == "高",
                Risk.status == "开放"
            ).all()
            if proj_risks:
                risk_rows = [(r.title, r.probability, r.impact, r.level or "-", r.mitigation or "-") for r in proj_risks]
                doc.add_paragraph("  高影响开放风险：").runs[0].italic = True
                _add_table(doc, ["风险标题", "概率", "影响", "等级", "应对措施"], risk_rows)
    else:
        doc.add_paragraph("✅ 当前所有项目均处于正常状态，无需重点关注。")

    # ─────────────────────────────────────────────
    # 第三部分：全局重大风险与问题
    # ─────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "三、全局重大风险与问题（高等级）", level=2)

    _add_heading(doc, "3.1 高严重等级未关闭问题 Top 10", level=3)
    top_issues = db.query(Issue).filter(
        Issue.status != "已关闭",
        Issue.severity == "高"
    ).order_by(Issue.due_date).limit(10).all()

    if top_issues:
        # 查询项目名称
        proj_map = {p.id: p.name for p in all_projects}
        i_rows = [
            (proj_map.get(i.project_id, "?"), i.title, i.severity, i.assignee or "-",
             _fmt_date(i.due_date), i.status)
            for i in top_issues
        ]
        _add_table(doc, ["所属项目", "问题标题", "严重等级", "负责人", "期望解决日期", "状态"], i_rows)
    else:
        doc.add_paragraph("✅ 当前无高严重等级未关闭问题。")

    doc.add_paragraph()
    _add_heading(doc, "3.2 高影响开放风险 Top 10", level=3)
    top_risks = db.query(Risk).filter(
        Risk.status == "开放",
        Risk.impact == "高"
    ).limit(10).all()

    if top_risks:
        proj_map = {p.id: p.name for p in all_projects}
        r_rows = [
            (proj_map.get(r.project_id, "?"), r.title, r.probability, r.impact,
             r.level or "-", r.assignee or "-", r.mitigation or "-")
            for r in top_risks
        ]
        _add_table(doc, ["所属项目", "风险标题", "概率", "影响", "等级", "负责人", "应对措施"], r_rows)
    else:
        doc.add_paragraph("✅ 当前无高影响开放风险。")

    # ─────────────────────────────────────────────
    # 第四部分：近期里程碑跟踪
    # ─────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "四、近期里程碑跟踪", level=2)
    proj_map = {p.id: p.name for p in all_projects}

    # 4.1 近期已完成（过去14天内）
    past_14 = today - timedelta(days=14)
    _add_heading(doc, "4.1 近期已达成里程碑（过去14天）", level=3)
    recent_done = db.query(Milestone).filter(
        Milestone.status == "已完成",
        Milestone.actual_date >= past_14,
        Milestone.actual_date <= today
    ).order_by(Milestone.actual_date.desc()).all()

    if recent_done:
        rd_rows = [(proj_map.get(m.project_id, "?"), m.name, _fmt_date(m.plan_date), _fmt_date(m.actual_date)) for m in recent_done]
        _add_table(doc, ["所属项目", "里程碑", "计划日期", "实际完成日期"], rd_rows)
    else:
        doc.add_paragraph("近期14天内无已完成里程碑。")

    # 4.2 近期计划达成（未来14天）
    future_14 = today + timedelta(days=14)
    doc.add_paragraph()
    _add_heading(doc, "4.2 近期计划里程碑（未来14天）", level=3)
    upcoming = db.query(Milestone).filter(
        Milestone.status.in_(["未开始", "进行中"]),
        Milestone.plan_date >= today,
        Milestone.plan_date <= future_14
    ).order_by(Milestone.plan_date).all()

    if upcoming:
        up_rows = [(proj_map.get(m.project_id, "?"), m.name, _fmt_date(m.plan_date), m.status) for m in upcoming]
        _add_table(doc, ["所属项目", "里程碑", "计划日期", "当前状态"], up_rows)
    else:
        doc.add_paragraph("未来14天内无计划中的里程碑。")

    # 4.3 已逾期未完成
    doc.add_paragraph()
    _add_heading(doc, "4.3 已逾期未达成里程碑（⚠️）", level=3)
    overdue = db.query(Milestone).filter(
        Milestone.status.in_(["未开始", "进行中", "延期"]),
        Milestone.plan_date < today
    ).order_by(Milestone.plan_date).all()

    if overdue:
        od_rows = [(proj_map.get(m.project_id, "?"), m.name, _fmt_date(m.plan_date), m.status) for m in overdue]
        _add_table(doc, ["所属项目", "里程碑", "计划日期", "状态"], od_rows)
        # 标红逾期里程碑行（最后一列）
    else:
        doc.add_paragraph("✅ 当前无逾期未达成的里程碑。")

    # ─────────────────────────────────────────────
    # 第五部分：项目状态一览表
    # ─────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "五、项目状态一览表", level=2)
    active_projects = [p for p in all_projects if p.status != "已完成"]
    if active_projects:
        p_rows = []
        for p in active_projects:
            open_issues_count = db.query(func.count(Issue.id)).filter(
                Issue.project_id == p.id, Issue.status != "已关闭"
            ).scalar() or 0
            open_risks_count = db.query(func.count(Risk.id)).filter(
                Risk.project_id == p.id, Risk.status == "开放"
            ).scalar() or 0
            used = db.query(func.sum(ManDay.days)).filter(ManDay.project_id == p.id).scalar() or 0
            budget_pct = f"{min(int(used / p.budget_mandays * 100), 100)}%" if p.budget_mandays else "N/A"
            p_rows.append((
                p.code, p.name, p.manager or "-", p.phase, p.status,
                _fmt_date(p.plan_end), open_issues_count, open_risks_count, budget_pct
            ))
        _add_table(doc, [
            "项目编号", "项目名称", "项目经理", "阶段", "状态",
            "计划结束", "未关闭问题", "开放风险", "人天消耗%"
        ], p_rows)
    else:
        doc.add_paragraph("当前无进行中的项目。")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────
# 7. Excel 项目组合数据汇总（Portfolio Overview Excel）
# ──────────────────────────────────────────────────────────
def generate_portfolio_excel(db: Session, report_date: Optional[date] = None) -> bytes:
    """
    生成覆盖所有项目的多 Sheet Excel 组合汇总报告：
      - Sheet 1: 项目状态一览
      - Sheet 2: 全局高严重问题清单
      - Sheet 3: 全局高影响风险清单
      - Sheet 4: 逾期里程碑
    """
    today = report_date or date.today()
    all_projects = db.query(Project).all()
    proj_map = {p.id: p.name for p in all_projects}

    wb = openpyxl.Workbook()

    # ── Sheet 1: 项目状态一览 ──────────────────────
    ws1 = wb.active
    ws1.title = "项目状态一览"
    ws1.cell(row=1, column=1, value=f"PMO 项目组合整体报告  生成日期：{_fmt_date(today)}")
    ws1.cell(row=1, column=1).font = Font(bold=True, size=14, color=COLOR_HEADER)

    headers1 = ["项目编号", "项目名称", "客户/甲方", "项目经理", "阶段", "状态",
                "计划开始", "计划结束", "预算人天", "已用人天", "人天消耗%", "未关闭问题", "开放风险"]
    for c, h in enumerate(headers1, 1):
        cell = ws1.cell(row=3, column=c, value=h)
        _excel_header_style(cell)

    for r, p in enumerate(all_projects, 4):
        used = db.query(func.sum(ManDay.days)).filter(ManDay.project_id == p.id).scalar() or 0
        open_i = db.query(func.count(Issue.id)).filter(Issue.project_id == p.id, Issue.status != "已关闭").scalar() or 0
        open_r = db.query(func.count(Risk.id)).filter(Risk.project_id == p.id, Risk.status == "开放").scalar() or 0
        pct = int(used / p.budget_mandays * 100) if p.budget_mandays else 0
        row_data = [p.code, p.name, p.client or "", p.manager or "", p.phase, p.status,
                    _fmt_date(p.plan_start), _fmt_date(p.plan_end),
                    p.budget_mandays, round(used, 1), f"{pct}%", open_i, open_r]
        for c, val in enumerate(row_data, 1):
            cell = ws1.cell(row=r, column=c, value=val if val is not None else "")
            _data_style(cell, align="center" if c > 6 else "left")
            if c == 6 and p.status in STATUS_COLOR:
                _color_cell(cell, STATUS_COLOR[p.status])
    _auto_col_width(ws1)

    # ── Sheet 2: 全局高严重问题 ────────────────────
    ws2 = wb.create_sheet("高严重等级问题")
    headers2 = ["所属项目", "问题标题", "描述", "严重等级", "来源", "负责人",
                "提出日期", "期望解决", "实际解决", "状态", "解决措施"]
    for c, h in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=c, value=h)
        _excel_header_style(cell)
    top_issues = db.query(Issue).filter(
        Issue.status != "已关闭", Issue.severity == "高"
    ).order_by(Issue.due_date).all()
    for r, i in enumerate(top_issues, 2):
        row_data = [proj_map.get(i.project_id, "?"), i.title, i.description, i.severity,
                    i.source, i.assignee, _fmt_date(i.raised_date), _fmt_date(i.due_date),
                    _fmt_date(i.resolved_date), i.status, i.resolution]
        for c, val in enumerate(row_data, 1):
            cell = ws2.cell(row=r, column=c, value=val or "")
            _data_style(cell)
            if c == 4:
                _color_cell(cell, COLOR_RED)
    _auto_col_width(ws2)

    # ── Sheet 3: 全局高影响风险 ────────────────────
    ws3 = wb.create_sheet("高影响风险")
    headers3 = ["所属项目", "风险标题", "描述", "概率", "影响", "风险等级", "负责人", "应对措施", "状态"]
    for c, h in enumerate(headers3, 1):
        cell = ws3.cell(row=1, column=c, value=h)
        _excel_header_style(cell)
    top_risks = db.query(Risk).filter(
        Risk.status == "开放", Risk.impact == "高"
    ).all()
    for r, risk in enumerate(top_risks, 2):
        row_data = [proj_map.get(risk.project_id, "?"), risk.title, risk.description,
                    risk.probability, risk.impact, risk.level, risk.assignee, risk.mitigation, risk.status]
        for c, val in enumerate(row_data, 1):
            cell = ws3.cell(row=r, column=c, value=val or "")
            _data_style(cell)
            if c == 6 and risk.level in RISK_COLOR:
                _color_cell(cell, RISK_COLOR[risk.level])
    _auto_col_width(ws3)

    # ── Sheet 4: 逾期里程碑 ───────────────────────
    ws4 = wb.create_sheet("逾期里程碑")
    headers4 = ["所属项目", "里程碑名称", "计划日期", "当前状态", "逾期天数", "备注"]
    for c, h in enumerate(headers4, 1):
        cell = ws4.cell(row=1, column=c, value=h)
        _excel_header_style(cell)
    overdue_ms = db.query(Milestone).filter(
        Milestone.status.in_(["未开始", "进行中", "延期"]),
        Milestone.plan_date < today
    ).order_by(Milestone.plan_date).all()
    for r, m in enumerate(overdue_ms, 2):
        overdue_days = (today - m.plan_date).days if m.plan_date else 0
        row_data = [proj_map.get(m.project_id, "?"), m.name, _fmt_date(m.plan_date),
                    m.status, overdue_days, m.description or ""]
        for c, val in enumerate(row_data, 1):
            cell = ws4.cell(row=r, column=c, value=val if val is not None else "")
            _data_style(cell)
            if c == 5 and overdue_days > 0:
                _color_cell(cell, COLOR_RED)
    _auto_col_width(ws4)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
